"""
PRAHARI Document Converter — Stage 2 pre-processing.

Converts heterogeneous bidder document formats into something Gemini can process:
  - PDF   → uploaded as-is via Gemini Files API (native multimodal)
  - DOCX  → text extracted with python-docx, saved as .txt, uploaded as plain text
  - DOC   → same path as DOCX (python-docx handles both if well-formed)
  - JPG / PNG / JPEG / WEBP / TIFF → uploaded as-is (Gemini vision natively handles images)

Returns a (gemini_path, mime_type) tuple for the caller to pass to genai.upload_file.
"""

import os
import tempfile
from pathlib import Path

ACCEPTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif"}

MIME_TYPES = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc":  "application/msword",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png":  "image/png",
    ".webp": "image/webp",
    ".tiff": "image/tiff",
    ".tif":  "image/tiff",
    ".txt":  "text/plain",
}


def is_accepted(filename: str) -> bool:
    return Path(filename).suffix.lower() in ACCEPTED_EXTENSIONS


def prepare_for_gemini(source_path: str, original_filename: str) -> tuple[str, str]:
    """
    Given a locally-saved file, return (path_to_upload, mime_type).
    For DOCX/DOC files, extracts text and writes a .txt temp file next to the source.
    The caller is responsible for deleting any returned temp file.
    """
    ext = Path(original_filename).suffix.lower()

    if ext in {".docx", ".doc"}:
        txt_path = source_path + ".extracted.txt"
        _extract_docx_text(source_path, txt_path)
        return txt_path, "text/plain"

    mime = MIME_TYPES.get(ext, "application/octet-stream")
    return source_path, mime


def detect_document_kind(filename: str) -> str:
    """
    Returns a human-readable document kind label used in logging and audit events.
      'pdf'   — portable document (may be digital or scanned; Gemini handles both)
      'word'  — DOCX/DOC; text extracted before Gemini upload
      'image' — JPG/PNG/WEBP/TIFF; processed via Gemini vision OCR
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".doc"}:
        return "word"
    if ext in {".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif"}:
        return "image"
    return "unknown"


def _extract_docx_text(docx_path: str, output_txt_path: str) -> None:
    """Extract all paragraph + table text from a .docx and write to a plain-text file."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise RuntimeError("python-docx is not installed. Add it to requirements.txt.")

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)

    full_text = "\n".join(lines)

    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(f"[Extracted from: {Path(docx_path).name}]\n\n")
        f.write(full_text)

    print(f"✓ DOCX→TXT: extracted {len(lines)} lines from {Path(docx_path).name}")
