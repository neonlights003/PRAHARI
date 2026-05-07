"""
Generate PRAHARI demo documents for the mock tender scenario.

Scenario: Supply and Installation of Bullet-Resistant Jackets — CRPF 2024-25
  NIT No: CRPF/HQ/PPE/2024-25/001

Run:
    cd /path/to/AI-For-Bharat
    python sample_data/generate_demo_docx.py

Outputs four .docx files in this directory:
    tender_nit_crpf_ppe_2024.docx      — Upload as the Tender NIT document
    bidder1_kavach_armour.docx         — Eligible bidder
    bidder2_frontier_defence.docx      — Manual Review (turnover borderline)
    bidder3_suraksha_equipment.docx    — Not Eligible (no GST registration)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value


# ─────────────────────────────────────────────────────────────────────────────
# 1. Tender NIT document
# ─────────────────────────────────────────────────────────────────────────────

def make_tender():
    doc = Document()

    heading(doc, "NOTICE INVITING TENDER (NIT)", level=1)
    heading(doc, "Central Reserve Police Force — Ministry of Home Affairs", level=2)
    para(doc, "")

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Field"
    t.rows[0].cells[1].text = "Details"
    data = [
        ("NIT No.",          "CRPF/HQ/PPE/2024-25/001"),
        ("Date of Issue",    "01 January 2025"),
        ("Tender Type",      "Open Competitive Tender"),
        ("Work Description", "Supply and Installation of Bullet-Resistant Jackets (NIJ Level IIIA) for CRPF Battalions — 2024-25"),
        ("Quantity",         "10,000 units"),
        ("Estimated Value",  "INR 25,00,00,000 (Twenty-Five Crore)"),
        ("Submission Deadline", "28 February 2025 at 15:00 hrs IST"),
        ("Bid Validity",     "180 days from submission date"),
        ("EMD Amount",       "INR 50,00,000 (Fifty Lakh)"),
        ("Contact",          "Director (Procurement), CRPF HQ, New Delhi — procurement@crpf.gov.in"),
    ]
    for label, value in data:
        table_row(t, label, value)

    para(doc, "")
    heading(doc, "SECTION A — SCOPE OF WORK", level=2)
    para(doc, (
        "The contractor shall supply, inspect, and deliver Bullet-Resistant Jackets (BRJs) conforming to "
        "NIJ Standard 0101.06 Level IIIA across designated CRPF battalion stores in 12 states. "
        "The jackets must provide full torso protection and integrate carrier vests compatible with existing CRPF load-bearing equipment."
    ))
    para(doc, "Delivery schedule: 3,000 units within 60 days; balance within 120 days of purchase order.")

    heading(doc, "SECTION B — ELIGIBILITY CRITERIA", level=2)
    para(doc, "Bidders must satisfy ALL mandatory criteria and are encouraged to meet preferred criteria.")
    para(doc, "")

    heading(doc, "B.1  Mandatory Criteria (disqualification if not met)", level=3)
    mandatory = [
        ("C001", "Annual Turnover",
         "The bidder's average annual turnover must be at least INR 10,00,00,000 (Ten Crore) "
         "over the last three financial years (FY 2021-22, 2022-23, 2023-24), as certified by a Chartered Accountant."),
        ("C002", "Prior Experience",
         "The bidder must have successfully completed a minimum of THREE (3) similar works "
         "(supply of body armour, ballistic protection, or PPE to defence/paramilitary forces) "
         "each valued at INR 1 crore or more, within the last five years. "
         "Completion certificates from the client must be enclosed."),
        ("C003", "GST Registration",
         "The bidder must possess a valid GST Registration Certificate in India. "
         "The GST number must be active and in good standing. Unregistered entities are not eligible."),
    ]
    for cid, name, desc in mandatory:
        heading(doc, f"[{cid}] {name} — MANDATORY", level=4)
        para(doc, desc)

    heading(doc, "B.2  Preferred Criteria (scored; non-compliance does not disqualify)", level=3)
    preferred = [
        ("C004", "ISO Certification",
         "ISO 9001:2015 certification from a NABCB-accredited body for manufacture of protective equipment or PPE. "
         "Bidders with this certification will receive additional weightage in technical scoring."),
        ("C005", "BIS/DGQA Approval",
         "Bureau of Indian Standards (BIS) licence or DGQA approval for the specific jacket model being offered. "
         "DGQA-approved suppliers will be given preference in case of equal technical scores."),
    ]
    for cid, name, desc in preferred:
        heading(doc, f"[{cid}] {name} — PREFERRED", level=4)
        para(doc, desc)

    heading(doc, "SECTION C — DOCUMENT CHECKLIST", level=2)
    checklist = [
        "Duly filled and signed Bid Form",
        "EMD in the form of Bank Guarantee from a scheduled bank",
        "Auditor-certified turnover statements for FY 2021-22, 2022-23, 2023-24",
        "Work completion certificates (minimum 3, for contracts ≥ INR 1 crore each)",
        "Valid GST Registration Certificate",
        "PAN Card copy",
        "ISO 9001:2015 Certificate (if applicable)",
        "BIS Licence or DGQA approval letter (if applicable)",
        "Manufacturer's authorisation letter (for non-manufacturers)",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    para(doc, "")
    para(doc, (
        "CRPF reserves the right to reject any or all bids without assigning reasons. "
        "This NIT does not constitute an offer or commitment to award a contract."
    ))

    doc.save(OUT / "tender_nit_crpf_ppe_2024.docx")
    print("✓ tender_nit_crpf_ppe_2024.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 2. Bidder 1 — Kavach Armour Solutions (Eligible)
# ─────────────────────────────────────────────────────────────────────────────

def make_bidder1():
    doc = Document()

    heading(doc, "TECHNICAL BID & QUALIFICATION DOCUMENTS", level=1)
    heading(doc, "Kavach Armour Solutions Pvt Ltd", level=2)
    para(doc, "Reference: NIT No. CRPF/HQ/PPE/2024-25/001")
    para(doc, "")

    heading(doc, "COMPANY PROFILE", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Field"
    t.rows[0].cells[1].text = "Details"
    company_data = [
        ("Company Name",      "Kavach Armour Solutions Pvt Ltd"),
        ("CIN",               "U25200DL2008PTC123456"),
        ("Registered Address","Plot 45, Industrial Area Phase II, Rohini, New Delhi — 110085"),
        ("GSTIN",             "07AABCK1234A1Z5"),
        ("PAN",               "AABCK1234A"),
        ("Contact Email",     "tenders@kavacharmour.in"),
        ("Contact Phone",     "+91-11-4567-8901"),
        ("Year of Incorporation", "2008"),
        ("Nature of Business", "Manufacturer & Supplier of Ballistic Protection Equipment"),
    ]
    for label, value in company_data:
        table_row(t, label, value)

    para(doc, "")
    heading(doc, "CRITERION C001 — ANNUAL TURNOVER", level=2)
    para(doc, (
        "Kavach Armour Solutions Pvt Ltd has consistently maintained turnover well above the threshold. "
        "The following figures are certified by M/s Sharma & Co., Chartered Accountants, New Delhi (ICAI Membership No. 056789)."
    ))
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Financial Year"
    t2.rows[0].cells[1].text = "Turnover (INR)"
    t2.rows[0].cells[2].text = "Turnover (in Crore)"
    for fy, amt, crore in [
        ("FY 2021-22", "16,20,00,000", "16.20"),
        ("FY 2022-23", "18,95,00,000", "18.95"),
        ("FY 2023-24", "20,30,00,000", "20.30"),
    ]:
        row = t2.add_row()
        row.cells[0].text = fy
        row.cells[1].text = amt
        row.cells[2].text = crore
    para(doc, "Average Annual Turnover: INR 18.48 Crore. This exceeds the mandatory threshold of INR 10 Crore.")
    para(doc, "CA Certificate reference: SHA/2024/CRPF/001, dated 15 December 2024.")

    para(doc, "")
    heading(doc, "CRITERION C002 — PRIOR EXPERIENCE", level=2)
    para(doc, "The following four (4) similar work completions are submitted in support of our eligibility:")
    works = [
        ("W-01", "CISF, MHA",     "Supply of 8,000 Body Armour Jackets (NIJ IIIA)",   "INR 2,15,00,000", "Completed March 2022", "Certificate enclosed as Annex-1A"),
        ("W-02", "BSF, MHA",      "Supply of 5,500 Ballistic Helmets and Vest Sets",   "INR 1,85,00,000", "Completed August 2022", "Certificate enclosed as Annex-1B"),
        ("W-03", "CRPF, MHA",     "Supply of 12,000 Stab-Resistant Vests",             "INR 3,10,00,000", "Completed January 2024", "Certificate enclosed as Annex-1C"),
        ("W-04", "SSB, MHA",      "Supply of 4,000 Multi-Threat Protection Jackets",   "INR 1,40,00,000", "Completed September 2024", "Certificate enclosed as Annex-1D"),
    ]
    for wid, client, desc, value, status, note in works:
        para(doc, f"[{wid}] {client}: {desc} | Value: {value} | {status} | {note}")

    para(doc, "")
    heading(doc, "CRITERION C003 — GST REGISTRATION", level=2)
    para(doc, "GST Registration Number: 07AABCK1234A1Z5")
    para(doc, "State: Delhi | Registration Status: Active | Valid from: 01 July 2017")
    para(doc, "GST Certificate enclosed as Annex-2.")

    para(doc, "")
    heading(doc, "CRITERION C004 — ISO CERTIFICATION (Preferred)", level=2)
    para(doc, "ISO 9001:2015 Certificate No: DNV-ISO-9001-2024-0456")
    para(doc, "Issued by: Det Norske Veritas (DNV), NABCB accredited | Scope: Design, Manufacture and Supply of Ballistic Protection Equipment")
    para(doc, "Valid until: 30 June 2027 | Certificate enclosed as Annex-3.")

    para(doc, "")
    heading(doc, "CRITERION C005 — BIS/DGQA APPROVAL (Preferred)", level=2)
    para(doc, "DGQA Approval Reference: DGQA/APP/BRJ/2023/0089")
    para(doc, "Approved model: KA-PRO-IIIA Bullet-Resistant Jacket | Date of approval: 22 March 2023")
    para(doc, "Approval certificate enclosed as Annex-4.")

    para(doc, "")
    para(doc, "We hereby declare that all information furnished above is true and correct to the best of our knowledge.")
    para(doc, "Authorised Signatory: Rajesh Kumar Verma | Designation: Managing Director")
    para(doc, "Date: 25 January 2025 | Place: New Delhi")

    doc.save(OUT / "bidder1_kavach_armour.docx")
    print("✓ bidder1_kavach_armour.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 3. Bidder 2 — Frontier Defence Gear Ltd (Manual Review — borderline turnover)
# ─────────────────────────────────────────────────────────────────────────────

def make_bidder2():
    doc = Document()

    heading(doc, "TECHNICAL BID & QUALIFICATION DOCUMENTS", level=1)
    heading(doc, "Frontier Defence Gear Ltd", level=2)
    para(doc, "Reference: NIT No. CRPF/HQ/PPE/2024-25/001")
    para(doc, "")

    heading(doc, "COMPANY PROFILE", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Field"
    t.rows[0].cells[1].text = "Details"
    company_data = [
        ("Company Name",      "Frontier Defence Gear Ltd"),
        ("CIN",               "U25200MH2012PLC234567"),
        ("Registered Address","Unit 7, MIDC Industrial Estate, Bhiwandi, Thane — 421302, Maharashtra"),
        ("GSTIN",             "27AABCF5678B1Z3"),
        ("PAN",               "AABCF5678B"),
        ("Contact Email",     "procurement@frontierdefence.com"),
        ("Contact Phone",     "+91-22-6789-0123"),
        ("Year of Incorporation", "2012"),
        ("Nature of Business", "Trader and Systems Integrator for Protective Equipment"),
    ]
    for label, value in company_data:
        table_row(t, label, value)

    para(doc, "")
    heading(doc, "CRITERION C001 — ANNUAL TURNOVER", level=2)
    para(doc, (
        "Certified by M/s Patel & Associates, Chartered Accountants, Mumbai (ICAI No. 078901). "
        "Note: FY 2021-22 turnover was impacted by Covid-19 supply chain disruptions."
    ))
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Financial Year"
    t2.rows[0].cells[1].text = "Turnover (INR)"
    t2.rows[0].cells[2].text = "Turnover (in Crore)"
    for fy, amt, crore in [
        ("FY 2021-22", "7,50,00,000",  "7.50"),
        ("FY 2022-23", "10,80,00,000", "10.80"),
        ("FY 2023-24", "11,20,00,000", "11.20"),
    ]:
        row = t2.add_row()
        row.cells[0].text = fy
        row.cells[1].text = amt
        row.cells[2].text = crore
    para(doc, "Average Annual Turnover: INR 9.83 Crore.")
    para(doc, (
        "We respectfully submit that our turnover for the most recent two financial years (FY 2022-23 and FY 2023-24) "
        "consistently exceeds the INR 10 Crore threshold. We request that the evaluation officer consider the upward trend "
        "and the Covid-19 impact on FY 2021-22 figures. Our order book as of January 2025 stands at INR 14.5 Crore."
    ))

    para(doc, "")
    heading(doc, "CRITERION C002 — PRIOR EXPERIENCE", level=2)
    para(doc, "Five (5) qualifying works completed:")
    works = [
        ("W-01", "NDRF, MHA",    "Supply of 3,000 Riot Control Protective Suits",        "INR 1,10,00,000", "Completed June 2020"),
        ("W-02", "Maharashtra Police", "Supply of 5,000 Anti-Stab Vests",                "INR 1,25,00,000", "Completed February 2021"),
        ("W-03", "ITBP, MHA",    "Supply of 2,500 Cold Weather Protective Jackets",       "INR 1,60,00,000", "Completed October 2022"),
        ("W-04", "BSF, MHA",     "Supply of 4,000 Ballistic Plate Carriers",              "INR 2,20,00,000", "Completed April 2023"),
        ("W-05", "CRPF, MHA",    "Supply of 3,500 Level IIIA Bullet-Resistant Jackets",  "INR 1,95,00,000", "Completed November 2024"),
    ]
    for wid, client, desc, value, status in works:
        para(doc, f"[{wid}] {client}: {desc} | Value: {value} | {status}")

    para(doc, "")
    heading(doc, "CRITERION C003 — GST REGISTRATION", level=2)
    para(doc, "GST Registration Number: 27AABCF5678B1Z3")
    para(doc, "State: Maharashtra | Registration Status: Active | Valid from: 15 August 2017")

    para(doc, "")
    heading(doc, "CRITERION C004 — ISO CERTIFICATION (Preferred)", level=2)
    para(doc, "Frontier Defence Gear Ltd does not currently hold ISO 9001:2015 certification.")
    para(doc, "We have initiated the certification process with Bureau Veritas and expect certification by Q3 2025.")

    para(doc, "")
    heading(doc, "CRITERION C005 — BIS/DGQA APPROVAL (Preferred)", level=2)
    para(doc, "DGQA Approval Reference: DGQA/APP/BRJ/2022/0045")
    para(doc, "Approved model: FD-PROTECT-IIIA Jacket | Date of approval: 10 July 2022")

    para(doc, "")
    para(doc, "Authorised Signatory: Sunil Mehta | Designation: Director, Operations")
    para(doc, "Date: 26 January 2025 | Place: Mumbai")

    doc.save(OUT / "bidder2_frontier_defence.docx")
    print("✓ bidder2_frontier_defence.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 4. Bidder 3 — Suraksha Equipment Suppliers (Not Eligible — no GST)
# ─────────────────────────────────────────────────────────────────────────────

def make_bidder3():
    doc = Document()

    heading(doc, "TECHNICAL BID & QUALIFICATION DOCUMENTS", level=1)
    heading(doc, "Suraksha Equipment Suppliers", level=2)
    para(doc, "Reference: NIT No. CRPF/HQ/PPE/2024-25/001")
    para(doc, "")

    heading(doc, "COMPANY PROFILE", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Field"
    t.rows[0].cells[1].text = "Details"
    company_data = [
        ("Company Name",      "Suraksha Equipment Suppliers"),
        ("Type",              "Proprietorship Firm"),
        ("Registered Address","No. 12, Kilpauk Industrial Estate, Chennai — 600010, Tamil Nadu"),
        ("GST Status",        "Composition Scheme — Surrendered registration, currently under re-registration process"),
        ("PAN",               "AABCS9012C"),
        ("Contact Email",     "suraksha.equip@gmail.com"),
        ("Contact Phone",     "+91-44-2345-6789"),
        ("Year of Establishment", "2005"),
        ("Nature of Business", "Manufacturer of Ballistic and Protective Gear"),
    ]
    for label, value in company_data:
        table_row(t, label, value)

    para(doc, "")
    heading(doc, "CRITERION C001 — ANNUAL TURNOVER", level=2)
    para(doc, "Certified by M/s Rajan & Co., Chartered Accountants, Chennai.")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Financial Year"
    t2.rows[0].cells[1].text = "Turnover (INR)"
    t2.rows[0].cells[2].text = "Turnover (in Crore)"
    for fy, amt, crore in [
        ("FY 2021-22", "19,80,00,000", "19.80"),
        ("FY 2022-23", "22,10,00,000", "22.10"),
        ("FY 2023-24", "24,50,00,000", "24.50"),
    ]:
        row = t2.add_row()
        row.cells[0].text = fy
        row.cells[1].text = amt
        row.cells[2].text = crore
    para(doc, "Average Annual Turnover: INR 22.13 Crore. Well above the mandatory threshold.")

    para(doc, "")
    heading(doc, "CRITERION C002 — PRIOR EXPERIENCE", level=2)
    para(doc, "Six (6) qualifying works completed in the last five years:")
    works = [
        ("W-01", "Tamil Nadu Police",  "Supply of 10,000 Anti-Riot Helmets and Shields", "INR 1,80,00,000", "Completed 2020"),
        ("W-02", "CRPF (Chennai Gp)", "Supply of 6,000 Level II Bullet-Resistant Vests", "INR 2,40,00,000", "Completed 2021"),
        ("W-03", "Karnataka Police",  "Supply of 8,000 Stab-Resistant Jackets",          "INR 2,10,00,000", "Completed 2022"),
        ("W-04", "CISF (South)",      "Supply of 5,000 Multi-Threat Vests",              "INR 1,90,00,000", "Completed 2022"),
        ("W-05", "NDRF",              "Supply of 3,000 Rescue Protective Suits",         "INR 1,20,00,000", "Completed 2023"),
        ("W-06", "Andhra Pradesh Spl Police", "Supply of 7,000 Level IIIA Jackets",     "INR 3,50,00,000", "Completed 2024"),
    ]
    for wid, client, desc, value, status in works:
        para(doc, f"[{wid}] {client}: {desc} | Value: {value} | {status}")

    para(doc, "")
    heading(doc, "CRITERION C003 — GST REGISTRATION", level=2)
    para(doc, (
        "Suraksha Equipment Suppliers was previously registered under the GST Composition Scheme "
        "(GSTIN: 33AABCS9012C1ZQ). Due to a change in business category, we surrendered the Composition Scheme "
        "registration in October 2024 and have applied for regular GST registration. "
        "The new registration application (ARN: AA3310245678901) is currently under processing by the GST department. "
        "We expect the new GSTIN to be issued within 30 working days."
    ))
    para(doc, "We request that our application may be considered favourably pending GST re-registration.")

    para(doc, "")
    heading(doc, "CRITERION C004 — ISO CERTIFICATION (Preferred)", level=2)
    para(doc, "ISO 9001:2015 Certificate No: TÜV-SÜD-ISO-9001-2023-4521")
    para(doc, "Issued by: TÜV SÜD South Asia | Valid until: 31 March 2026")

    para(doc, "")
    heading(doc, "CRITERION C005 — BIS/DGQA APPROVAL (Preferred)", level=2)
    para(doc, "DGQA approval application submitted in November 2024 for model SE-GUARD-IIIA. Approval pending.")

    para(doc, "")
    para(doc, "Authorised Signatory: P. Venkataraman | Designation: Proprietor")
    para(doc, "Date: 27 January 2025 | Place: Chennai")

    doc.save(OUT / "bidder3_suraksha_equipment.docx")
    print("✓ bidder3_suraksha_equipment.docx")


if __name__ == "__main__":
    make_tender()
    make_bidder1()
    make_bidder2()
    make_bidder3()
    print("\nAll demo documents generated in sample_data/")
    print("See sample_data/README_DEMO.md for the upload walkthrough.")
