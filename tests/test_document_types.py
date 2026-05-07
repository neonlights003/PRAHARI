"""
Tests for PRAHARI document format support.
Verifies that the document converter accepts all required formats
(typed PDFs, scanned PDFs, Word files, photographs) as mandated by the
problem statement non-negotiable: "The system must handle scanned documents
and photographs, not only digital text."
"""

import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from backend.document_converter import is_accepted, detect_document_kind, prepare_for_gemini
import tempfile
import pathlib


# ── Acceptance gate ───────────────────────────────────────────────────────────

ACCEPTED_FORMATS = [
    "balance_sheet.pdf",
    "work_order.PDF",           # uppercase extension
    "certificate.docx",
    "old_format.doc",
    "scanned_certificate.jpg",
    "photo_of_registration.jpeg",
    "screenshot.png",
    "high_res_scan.tiff",
    "modern_scan.tif",
    "compressed.webp",
]

REJECTED_FORMATS = [
    "malware.exe",
    "data.csv",
    "spreadsheet.xlsx",
    "presentation.pptx",
    "archive.zip",
    "script.py",
    "",
]


def test_all_required_formats_accepted():
    for filename in ACCEPTED_FORMATS:
        assert is_accepted(filename), f"Expected {filename!r} to be accepted"


def test_rejected_formats_blocked():
    for filename in REJECTED_FORMATS:
        assert not is_accepted(filename), f"Expected {filename!r} to be rejected"


# ── Document kind detection ───────────────────────────────────────────────────

def test_pdf_kind():
    assert detect_document_kind("report.pdf") == "pdf"
    assert detect_document_kind("TENDER.PDF") == "pdf"


def test_word_kind():
    assert detect_document_kind("proposal.docx") == "word"
    assert detect_document_kind("old.doc") == "word"


def test_image_kind():
    for name in ["cert.jpg", "cert.jpeg", "cert.png", "cert.tiff", "cert.tif", "cert.webp"]:
        assert detect_document_kind(name) == "image", f"Expected image for {name!r}"


# ── DOCX text extraction ──────────────────────────────────────────────────────

def test_docx_extraction_creates_txt_file():
    """
    Creates a minimal in-memory .docx, writes it to a temp file,
    runs prepare_for_gemini, and confirms a .txt file is produced.
    """
    try:
        from docx import Document
    except ImportError:
        import pytest
        pytest.skip("python-docx not installed")

    # Build a minimal DOCX in memory
    doc = Document()
    doc.add_paragraph("Annual Turnover: INR 7.5 Crore (FY 2022-23)")
    doc.add_paragraph("GST Registration Number: 27AABCU9603R1ZX")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        out_path, mime = prepare_for_gemini(tmp_path, "submission.docx")
        assert mime == "text/plain"
        assert out_path.endswith(".txt")
        assert pathlib.Path(out_path).exists()

        content = pathlib.Path(out_path).read_text(encoding="utf-8")
        assert "Annual Turnover" in content
        assert "GST Registration" in content
    finally:
        pathlib.Path(tmp_path).unlink(missing_ok=True)
        if out_path != tmp_path:
            pathlib.Path(out_path).unlink(missing_ok=True)


def test_pdf_passthrough():
    """PDF files are returned as-is — no conversion."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(b"%PDF-1.4 fake content")
        tmp_path = tmp.name
    try:
        out_path, mime = prepare_for_gemini(tmp_path, "scan.pdf")
        assert out_path == tmp_path
        assert mime == "application/pdf"
    finally:
        pathlib.Path(tmp_path).unlink(missing_ok=True)


def test_image_passthrough():
    """Image files are returned as-is with the correct MIME type."""
    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
        tmp.write(b"\xff\xd8\xff fake jpeg")
        tmp_path = tmp.name
    try:
        out_path, mime = prepare_for_gemini(tmp_path, "certificate_photo.jpg")
        assert out_path == tmp_path
        assert mime == "image/jpeg"
    finally:
        pathlib.Path(tmp_path).unlink(missing_ok=True)
